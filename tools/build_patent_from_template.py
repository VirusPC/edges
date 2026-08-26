#!/usr/bin/env python3
"""Fill the official patent disclosure template with REDoc Markdown content."""

from __future__ import annotations

import argparse
import html
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


BODY_FONT = "宋体"
HEADING_FONT = "宋体"
MONO_FONT = "Consolas"
INK = "000000"
MUTED = "555555"
LINK_BLUE = "0563C1"
TABLE_HEADER = "E7E6E6"
CODE_BG = "F2F2F2"
BODY_WIDTH_DXA = 8280


SECTION_TITLES = {
    "1": "0 缩略语和关键术语定义",
    "2": "1 本发明的应用场景说明",
    "3": "2 详细介绍技术背景，并描述本发明的目的及解决的技术问题",
    "4": "3 本发明的技术方案具体是怎么做的",
    "5": "4 本发明的亮点/技术关键点/想要保护的点是什么，以及对应的好处",
}


def set_font(run, *, size=10.5, bold=False, italic=False, color=INK, font=BODY_FONT) -> None:
    run.font.name = font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_fonts.set(qn("w:eastAsia"), BODY_FONT if font == MONO_FONT else font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def configure_paragraph(paragraph, *, before=0, after=6, line=1.15, keep=False) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.widow_control = True
    pf.keep_with_next = keep


def add_linked_page_section(doc: Document) -> None:
    previous = doc.sections[-1]
    header_children = [deepcopy(child) for child in previous.header._element]
    footer_children = [deepcopy(child) for child in previous.footer._element]
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = previous.page_width
    section.page_height = previous.page_height
    section.left_margin = previous.left_margin
    section.right_margin = previous.right_margin
    section.top_margin = previous.top_margin
    section.bottom_margin = previous.bottom_margin
    section.header_distance = previous.header_distance
    section.footer_distance = previous.footer_distance
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for container, children in (
        (section.header._element, header_children),
        (section.footer._element, footer_children),
    ):
        for child in list(container):
            container.remove(child)
        for child in children:
            container.append(child)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), BODY_FONT)
    r_fonts.set(qn("w:hAnsi"), BODY_FONT)
    r_fonts.set(qn("w:eastAsia"), BODY_FONT)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "21")
    r_pr.extend([r_fonts, color, underline, size])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\))")


def clean_markup(text: str) -> str:
    text = re.sub(r"</?font[^>]*>", "", text)
    return html.unescape(text.replace("<br/>", "\n").replace("</br>", "\n"))


def add_inline(paragraph, raw: str, *, size=10.5, color=INK) -> None:
    raw = clean_markup(raw)
    pos = 0
    for match in INLINE_RE.finditer(raw):
        if match.start() > pos:
            set_font(paragraph.add_run(raw[pos:match.start()]), size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            set_font(paragraph.add_run(token[2:-2]), size=size, bold=True, color=color)
        elif token.startswith("`"):
            set_font(paragraph.add_run(token[1:-1]), size=size - 0.5, color=color, font=MONO_FONT)
        else:
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link:
                add_hyperlink(paragraph, link.group(1), link.group(2))
        pos = match.end()
    if pos < len(raw):
        set_font(paragraph.add_run(raw[pos:]), size=size, color=color)


def set_cell_text(cell, text: str, *, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    configure_paragraph(p, after=0, line=1.05)
    set_font(p.add_run(text), size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        hdr = OxmlElement("w:tblHeader")
        hdr.set(qn("w:val"), "true")
        tr_pr.append(hdr)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(idx, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    if cols == 3:
        return [1900, 1900, BODY_WIDTH_DXA - 3800]
    if cols == 2:
        return [2500, BODY_WIDTH_DXA - 2500]
    scores = []
    for idx in range(cols):
        values = [min(max(len(re.sub(r"\*\*|`", "", row[idx])), 6), 80) for row in rows]
        scores.append(max(sum(values) / max(len(values), 1), 8))
    total = sum(scores)
    widths = [max(900, round(BODY_WIDTH_DXA * score / total)) for score in scores]
    widths[-1] += BODY_WIDTH_DXA - sum(widths)
    return widths


def parse_table_row(line: str) -> list[str]:
    return [clean_markup(cell.strip()) for cell in re.split(r"(?<!\\)\|", line.strip().strip("|"))]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    if len(rows) > 10:
        header, body_rows = rows[0], rows[1:]
        for part, start in enumerate(range(0, len(body_rows), 3)):
            if part:
                continuation = doc.add_paragraph(style="Normal")
                configure_paragraph(continuation, after=5, line=1.0, keep=True)
                set_font(continuation.add_run("（续表）"), size=10, bold=True, color=MUTED)
            add_table(doc, [header] + body_rows[start : start + 3])
        return
    cols = max(len(row) for row in rows)
    rows = [row + [""] * (cols - len(row)) for row in rows]
    table = doc.add_table(rows=len(rows), cols=cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        table.style = "Normal Table"
    set_table_geometry(table, table_widths(rows))
    set_repeat_table_header(table.rows[0])
    for row_idx, row in enumerate(rows):
        prevent_row_split(table.rows[row_idx])
        for col_idx, text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            if row_idx == 0:
                set_cell_shading(cell, TABLE_HEADER)
            cell.text = ""
            p = cell.paragraphs[0]
            configure_paragraph(p, after=1, line=1.05)
            p.paragraph_format.keep_together = True
            if row_idx < len(rows) - 1:
                p.paragraph_format.keep_with_next = True
            add_inline(p, text, size=9.0 if cols >= 3 else 9.5)
            if row_idx == 0:
                for run in p.runs:
                    run.bold = True
    spacer = doc.add_paragraph()
    configure_paragraph(spacer, after=2, line=1.0)


def add_glossary(doc: Document, rows: list[list[str]]) -> None:
    """Render the long terminology table as compact definition entries.

    This preserves all three source fields while allowing Word to paginate the
    glossary naturally without a multi-page table obscuring the template header.
    """
    for row in rows[1:]:
        term = row[0] if row else ""
        english = row[1] if len(row) > 1 else ""
        definition = row[2] if len(row) > 2 else ""
        p = doc.add_paragraph(style="Normal")
        configure_paragraph(p, after=8, line=1.12)
        p.paragraph_format.keep_together = True
        p.paragraph_format.left_indent = Inches(0.08)
        label = term if not english or english == "—" else f"{term}（{english}）"
        set_font(p.add_run(label), size=10.5, bold=True)
        p.add_run().add_break()
        add_inline(p, definition, size=10.2)


def add_section_heading(doc: Document, text: str, *, tips=False) -> None:
    p = doc.add_paragraph(style="Normal (Web)" if "Normal (Web)" in [s.name for s in doc.styles] else "Normal")
    configure_paragraph(p, before=14, after=8, line=1.05, keep=True)
    p.paragraph_format.keep_together = True
    set_font(p.add_run(text), size=16, bold=False, font=HEADING_FONT)


def add_subheading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    configure_paragraph(p, before=10, after=5, line=1.05, keep=True)
    p.paragraph_format.keep_together = True
    set_font(p.add_run(text), size=12, bold=True, font=HEADING_FONT)


def add_body_paragraph(doc: Document, text: str, *, left_indent=0.0, first_line=0.0) -> None:
    # The source document numbers its implementation section as chapter 4, while
    # the official template places that material in chapter 3.
    text = re.sub(r"^(\*\*图 )4\.", r"\g<1>3.", text)
    p = doc.add_paragraph(style="Normal")
    configure_paragraph(p, after=6, line=1.15)
    if re.sub(r"\*\*", "", text).strip().startswith("图 "):
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.keep_together = True
    if left_indent:
        p.paragraph_format.left_indent = Inches(left_indent)
    if first_line:
        p.paragraph_format.first_line_indent = Inches(first_line)
    add_inline(p, text)


def add_quote(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph(style="Normal")
    configure_paragraph(p, before=3, after=7, line=1.12)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.1)
    shade_paragraph(p, "F2F2F2")
    add_inline(p, "\n".join(lines), size=10, color=MUTED)


def add_code(doc: Document, language: str, code_lines: list[str]) -> None:
    if language:
        p = doc.add_paragraph(style="Normal")
        configure_paragraph(p, before=4, after=2, line=1.0, keep=True)
        set_font(p.add_run(language.upper()), size=8.5, bold=True, color=MUTED)
    for line in code_lines or [""]:
        p = doc.add_paragraph(style="Normal")
        configure_paragraph(p, after=0, line=1.0)
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.right_indent = Inches(0.12)
        shade_paragraph(p, CODE_BG)
        set_font(p.add_run(line or " "), size=8.3, font=MONO_FONT)
    spacer = doc.add_paragraph()
    configure_paragraph(spacer, after=4, line=1.0)


def add_diagram(doc: Document, png: Path, number: int) -> None:
    if not png.exists():
        add_body_paragraph(doc, f"[图 {number} 未能渲染]", left_indent=0.2)
        return
    with Image.open(png) as image:
        ratio = image.height / image.width
    width = 5.55
    height = width * ratio
    if height > 8.35:
        height = 8.35
        width = height / ratio
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    configure_paragraph(p, before=5, after=7, line=1.0)
    p.paragraph_format.keep_together = True
    descriptions = {
        1: "本发明总体架构与知识生命周期闭环图",
        2: "双层存储与加载流程图",
        3: "入口索引与渐进式披露检索流程图",
        4: "主动与被动触发的记忆沉淀流程图",
        5: "预处理建索引与搜索子智能体检索流程图",
        6: "两条积累路径与 Dream 五阶段流水线图",
        7: "团队记忆防腐巡检流程图",
    }
    picture = p.add_run().add_picture(str(png), width=Inches(width), height=Inches(height))
    picture._inline.docPr.set("title", f"图 {number}")
    picture._inline.docPr.set("descr", descriptions[number])


def prepare_template(doc: Document) -> None:
    body = doc._element.body
    title_node = None
    table_node = None
    for child in list(body):
        tag = child.tag.split("}")[-1]
        text = "".join(child.itertext())
        if tag == "p" and "技术交底书撰写模板" in text and title_node is None:
            title_node = child
        elif tag == "tbl" and table_node is None:
            table_node = child
    if title_node is None or table_node is None:
        raise RuntimeError("Unable to locate the title and metadata table in the template")
    for child in list(body):
        if child not in (title_node, table_node) and child.tag != qn("w:sectPr"):
            body.remove(child)

    title_p = next(p for p in doc.paragraphs if p._p is title_node)
    clear_paragraph(title_p)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    configure_paragraph(title_p, before=12, after=16, line=1.0, keep=True)
    set_font(title_p.add_run("技术交底书"), size=22, bold=True, font=HEADING_FONT)

    table = doc.tables[0]
    set_repeat_table_header(table.rows[0])
    set_cell_text(table.cell(0, 0), "交底书名称（即发明名称）", bold=True, size=10)
    set_cell_text(
        table.cell(1, 0),
        "一种面向团队协作场景的基于文件系统的 Agent 记忆检索、沉淀与防腐方法",
        size=10.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_cell_text(table.cell(2, 0), "发明人（递交以 OA 单据填写的发明人为准）", bold=True, size=9.5)
    set_cell_text(table.cell(2, 1), "所属部门", bold=True, size=9.5)
    set_cell_text(table.cell(3, 0), "守岁", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(3, 1), "业务技术部/社区工程部/活动运营 PE 组", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(4, 0), "交底书撰写人（即技术联系人）", bold=True, size=9.5)
    set_cell_text(table.cell(4, 1), "联系人电话及工作邮箱", bold=True, size=9.5)
    set_cell_text(table.cell(5, 0), "守岁", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(5, 1), "15778001136；pengcheng1@xiaohongshu.com", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_front_note(doc: Document) -> None:
    p = doc.add_paragraph(style="Normal")
    configure_paragraph(p, before=8, after=8, line=1.12)
    shade_paragraph(p, "F2F2F2")
    note = (
        "说明：本交底书基于项目内部文档《AI Memory · 技术规划》"
        "《基于长期记忆的团队级上下文管理》《AI-Memory 架构图》"
        "《dream-server 架构》整理填写。全文按“六个技术支柱”组织，"
        "对同一事物采用统一名称，外文均给出中文解释。"
    )
    set_font(p.add_run(note), size=9.5, color=MUTED)


def transform_subheading(text: str) -> str:
    text = re.sub(r"^3\.", "2.", text)
    text = re.sub(r"^4\.", "3.", text)
    return text


def add_source_content(doc: Document, source: Path, diagrams_dir: Path) -> None:
    text = source.read_text(encoding="utf-8")
    text = re.sub(
        r'<redoc-text-draw remoteTemplate=".*?"\s+(?:remoteView="[^"]+"\s+)?theme="light"/>',
        lambda m, c=iter(range(1, 100)): f"\n[[DIAGRAM:{next(c)}]]\n",
        text,
        flags=re.S,
    )
    lines = text.splitlines()
    start = next(i for i, line in enumerate(lines) if re.match(r"^##\s+1、", line.strip()))
    lines = lines[start:]

    idx = 0
    first_section = True
    current_section = ""
    while idx < len(lines):
        stripped = lines[idx].strip()
        if not stripped or stripped == "<br/>":
            idx += 1
            continue
        if stripped.startswith("[[DIAGRAM:"):
            number = int(re.search(r"\d+", stripped).group())
            add_diagram(doc, diagrams_dir / f"diagram-{number}.png", number)
            idx += 1
            continue
        if stripped.startswith("```"):
            language = stripped[3:].strip()
            code_lines = []
            idx += 1
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code_lines.append(lines[idx])
                idx += 1
            add_code(doc, language, code_lines)
            idx += 1
            continue
        if stripped.startswith("|"):
            raw_rows = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                raw_rows.append(parse_table_row(lines[idx]))
                idx += 1
            rows = [
                row
                for row in raw_rows
                if not all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in row)
            ]
            if current_section == "1":
                add_glossary(doc, rows)
            else:
                add_table(doc, rows)
            continue
        top_heading = re.match(r"^##\s+([1-5])、(.*)$", stripped)
        if top_heading:
            if first_section:
                add_linked_page_section(doc)
            current_section = top_heading.group(1)
            add_section_heading(doc, SECTION_TITLES[top_heading.group(1)])
            if first_section:
                first_section = False
            idx += 1
            continue
        if re.match(r"^##\s+TIPS", stripped, flags=re.I):
            add_section_heading(doc, "TIPS：其他有助于专利代理人理解本发明的参考资料", tips=True)
            idx += 1
            continue
        subheading = re.match(r"^###\s+(.*)$", stripped)
        if subheading:
            add_subheading(doc, transform_subheading(subheading.group(1)))
            idx += 1
            continue
        if stripped.startswith(">"):
            quote_lines = []
            while idx < len(lines) and lines[idx].strip().startswith(">"):
                quote_lines.append(lines[idx].strip()[1:].strip())
                idx += 1
            add_quote(doc, quote_lines)
            continue
        bullet = re.match(r"^-\s+(.*)$", stripped)
        if bullet:
            add_body_paragraph(doc, "• " + bullet.group(1), left_indent=0.22)
            idx += 1
            continue
        number = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if number:
            add_body_paragraph(doc, f"{number.group(1)}. {number.group(2)}", left_indent=0.22)
            idx += 1
            continue
        cleaned = clean_markup(stripped).replace("【END】", "").strip()
        if cleaned:
            add_body_paragraph(doc, cleaned)
        idx += 1


def enable_field_updates(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def materialize_even_page_headers(doc: Document) -> None:
    """Give LibreOffice explicit left-page header/footer parts.

    The source template relies on Word's default header inheritance. Some
    LibreOffice builds otherwise paint full left-page body frames over the
    inherited header/footer, so explicit even-page parts are safer for QA.
    """
    doc.settings.odd_and_even_pages_header_footer = True
    for section in doc.sections:
        pairs = (
            (section.header, section.even_page_header),
            (section.footer, section.even_page_footer),
        )
        for source, target in pairs:
            children = [deepcopy(child) for child in source._element]
            target.is_linked_to_previous = False
            for child in list(target._element):
                target._element.remove(child)
            for child in children:
                target._element.append(child)


def build(template: Path, source: Path, diagrams_dir: Path, output: Path) -> None:
    doc = Document(template)
    prepare_template(doc)
    add_front_note(doc)
    add_source_content(doc, source, diagrams_dir)
    enable_field_updates(doc)
    materialize_even_page_headers(doc)
    doc.core_properties.title = "技术交底书：一种面向团队协作场景的基于文件系统的 Agent 记忆检索、沉淀与防腐方法"
    doc.core_properties.subject = "发明专利技术交底书"
    doc.core_properties.author = "守岁"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("template", type=Path)
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--diagrams-dir", type=Path, required=True)
    args = parser.parse_args()
    build(args.template, args.source, args.diagrams_dir, args.output)


if __name__ == "__main__":
    main()
