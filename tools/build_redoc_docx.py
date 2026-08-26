#!/usr/bin/env python3
"""Convert the retrieved REDoc-flavored Markdown patent disclosure to DOCX."""

from __future__ import annotations

import argparse
import html
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "202124"
MUTED = "5F6368"
LIGHT = "F2F4F7"
CALLOUT = "F4F6F9"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}
CJK_FONT = "Noto Sans CJK SC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGINS.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, *, size=None, bold=None, italic=None, color=INK, mono=False) -> None:
    latin = CJK_FONT
    east = CJK_FONT
    run.font.name = latin
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill: str, left_border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if left_border:
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), left_border)
        p_bdr.append(left)
        p_pr.append(p_bdr)


def add_hyperlink(paragraph, text: str, url: str, *, bold=False) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    if bold:
        b = OxmlElement("w:b")
        r_pr.append(b)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), CJK_FONT)
    r_pr.append(fonts)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\))")


def add_inline(paragraph, raw: str, *, base_size=11, base_color=INK) -> None:
    raw = html.unescape(re.sub(r"</?font[^>]*>", "", raw)).replace("<br/>", "\n")
    pos = 0
    for match in INLINE_RE.finditer(raw):
        if match.start() > pos:
            run = paragraph.add_run(raw[pos:match.start()])
            set_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=base_size, bold=True, color=base_color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=max(base_size - 0.5, 8), color=DARK_BLUE, mono=True)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "EDF2F7")
            run._element.get_or_add_rPr().append(shd)
        else:
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link:
                add_hyperlink(paragraph, link.group(1), link.group(2))
        pos = match.end()
    if pos < len(raw):
        run = paragraph.add_run(raw[pos:])
        set_font(run, size=base_size, color=base_color)


def clean_cell(text: str) -> str:
    return html.unescape(re.sub(r"</?font[^>]*>", "", text).replace("<br/>", "\n").strip())


def parse_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [clean_cell(part) for part in re.split(r"(?<!\\)\|", stripped)]


def table_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    scores = []
    for idx in range(cols):
        lengths = [min(max(len(re.sub(r"\*\*|`", "", row[idx])), 4), 90) for row in rows if idx < len(row)]
        score = max(sum(lengths) / max(len(lengths), 1), 8)
        scores.append(score)
    total = sum(scores)
    widths = [max(950, round(TABLE_WIDTH_DXA * score / total)) for score in scores]
    diff = TABLE_WIDTH_DXA - sum(widths)
    widths[-1] += diff
    if widths[-1] < 950:
        borrow = 950 - widths[-1]
        donor = max(range(cols - 1), key=lambda i: widths[i])
        widths[donor] -= borrow
        widths[-1] = 950
    return widths


def add_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    rows = [row + [""] * (cols - len(row)) for row in rows]
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    widths = table_widths(rows)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.08
            add_inline(p, text, base_size=9.2, base_color=INK)
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run("第 ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_font(run, size=9, color=MUTED)


def add_numbering(paragraph, doc: Document, start: int) -> None:
    """Attach a real decimal numbering definition starting at the source value."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start_el = OxmlElement("w:start")
    start_el.set(qn("w:val"), str(start))
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    lvl.extend([start_el, num_fmt, lvl_text, suffix, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    para_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    para_pr.append(num_pr)
def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = CJK_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = CJK_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = CJK_FONT
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = CJK_FONT
    code.font.size = Pt(8.5)
    code._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    code.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.15)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.line_spacing = 1.0


def preprocess(text: str, diagram_dir: Path) -> tuple[list[str], list[Path]]:
    diagram_dir.mkdir(parents=True, exist_ok=True)
    diagrams = []
    pattern = re.compile(
        r'<redoc-text-draw remoteTemplate="(.*?)"\s+(?:remoteView="[^"]+"\s+)?theme="light"/>',
        re.S,
    )

    def replace(match):
        idx = len(diagrams) + 1
        mermaid = html.unescape(match.group(1)).replace("</br>", "<br/>")
        path = diagram_dir / f"diagram-{idx}.mmd"
        path.write_text(mermaid + "\n", encoding="utf-8")
        diagrams.append(path)
        return f"\n[[DIAGRAM:{idx}]]\n"

    text = pattern.sub(replace, text)
    return text.splitlines(), diagrams


def build(source: Path, output: Path, diagrams_dir: Path, title: str) -> None:
    lines, diagrams = preprocess(source.read_text(encoding="utf-8"), diagrams_dir)
    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("发明专利技术交底书")
    set_font(hr, size=9, bold=True, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(fp)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(6)
    kr = kicker.add_run("发明专利技术交底书")
    set_font(kr, size=10, bold=True, color=BLUE)
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(16)
    title_p.paragraph_format.keep_with_next = True
    tr = title_p.add_run(title.replace("发明专利技术交底书：", ""))
    set_font(tr, size=21, bold=True, color=DARK_BLUE)

    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()
        if not stripped or stripped == "<br/>":
            idx += 1
            continue
        if stripped.startswith("[[DIAGRAM:"):
            number = int(re.search(r"\d+", stripped).group())
            png = diagrams_dir / f"diagram-{number}.png"
            if png.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run()
                with Image.open(png) as image:
                    ratio = image.height / image.width
                width = 6.25
                height = width * ratio
                if height > 8.45:
                    height = 8.45
                    width = height / ratio
                run.add_picture(str(png), width=Inches(width), height=Inches(height))
            else:
                p = doc.add_paragraph()
                add_inline(p, f"[图 {number} 未能渲染]", base_size=10, base_color=MUTED)
            idx += 1
            continue
        if stripped.startswith("```"):
            language = stripped[3:].strip()
            code_lines = []
            idx += 1
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code_lines.append(lines[idx])
                idx += 1
            if language:
                label = doc.add_paragraph()
                label.paragraph_format.space_before = Pt(5)
                label.paragraph_format.space_after = Pt(2)
                lr = label.add_run(language.upper())
                set_font(lr, size=8, bold=True, color=MUTED)
            for code_line in code_lines or [""]:
                p = doc.add_paragraph(style="Code Block")
                shade_paragraph(p, "F7F8FA")
                run = p.add_run(code_line or " ")
                set_font(run, size=8.5, color=INK, mono=True)
            idx += 1
            continue
        if stripped.startswith("|"):
            raw_rows = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                raw_rows.append(parse_table_row(lines[idx]))
                idx += 1
            rows = [row for row in raw_rows if not all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in row)]
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 3)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, heading.group(2), base_size={1: 16, 2: 13, 3: 12}[level], base_color=BLUE if level < 3 else DARK_BLUE)
            idx += 1
            continue
        if stripped.startswith(">"):
            quote_lines = []
            while idx < len(lines) and lines[idx].strip().startswith(">"):
                quote_lines.append(lines[idx].strip()[1:].strip())
                idx += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.16)
            p.paragraph_format.right_indent = Inches(0.08)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            shade_paragraph(p, CALLOUT, BLUE)
            add_inline(p, "\n".join(quote_lines), base_size=10, base_color=INK)
            continue
        bullet = re.match(r"^-\s+(.*)$", stripped)
        number = re.match(r"^(\d+)[.)]\s+(.*)$", stripped)
        if bullet or number:
            p = doc.add_paragraph(style="List Bullet" if bullet else None)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if number:
                add_numbering(p, doc, int(number.group(1)))
                add_inline(p, number.group(2))
            else:
                add_inline(p, bullet.group(1))
            idx += 1
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.widow_control = True
        add_inline(p, stripped)
        idx += 1

    doc.core_properties.title = title
    doc.core_properties.subject = "发明专利技术交底书"
    doc.core_properties.author = "守岁"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--diagrams-dir", type=Path, required=True)
    parser.add_argument("--title", required=True)
    args = parser.parse_args()
    build(args.source, args.output, args.diagrams_dir, args.title)


if __name__ == "__main__":
    main()
